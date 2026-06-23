import re
from flask import Flask, request, jsonify, send_from_directory, send_file, after_this_request, render_template
import os
import pandas as pd
from docx import Document
import uuid
import io
import zipfile
import requests
from datetime import datetime
import numpy as np
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
import json
import time
from threading import Thread
from datetime import datetime, timedelta
from flask import Flask, render_template, request, redirect, url_for, session, flash
from flask_mysqldb import MySQL
from flask_mail import Mail, Message
import bcrypt
import random
import string
from datetime import datetime, timedelta

from flask_login import LoginManager, login_user, login_required, logout_user, UserMixin, current_user
import secrets

from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from dotenv import load_dotenv
load_dotenv()

app = Flask(__name__, static_folder='static', template_folder='templates')
app.secret_key = 'your_secret_key'
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["200000 per day", "90000 per hour"]  # global default (optional)
)

# MySQL config
app.secret_key = os.environ.get('SECRET_KEY')
app.config['MYSQL_HOST'] = os.environ.get('MYSQL_HOST')
app.config['MYSQL_USER'] = os.environ.get('MYSQL_USER')
app.config['MYSQL_PASSWORD'] = os.environ.get('MYSQL_PASSWORD')
app.config['MYSQL_DB'] = os.environ.get('MYSQL_DB')
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD')
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USE_SSL'] = False

ADMIN_EMAIL = os.environ.get('ADMIN_EMAIL')

mysql = MySQL(app)
mail = Mail(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login_page' # or 'login' if that's your login route

class User(UserMixin):
    def __init__(self, id, email):
        self.id = id
        self.email = email

    def get_id(self):
        return str(self.id)

@login_manager.user_loader
def load_user(user_id):
    cur = mysql.connection.cursor()
    cur.execute("SELECT id, email FROM users WHERE id=%s", (user_id,))
    user = cur.fetchone()
    if user:
        return User(user[0], user[1])
    return None


def generate_otp(length=6):
    return ''.join(random.choices(string.digits, k=length))

@app.route('/', methods=['GET', 'POST'])
def login_page():
    if current_user.is_authenticated:
        return render_template('index.html', email=current_user.email)
    return render_template('login.html')


# LOGIN
@app.route('/login', methods=['POST'])
@limiter.limit("20 per minute") 
def login():
    email = request.form['email']
    password = request.form['password'].encode('utf-8')
    cur = mysql.connection.cursor()
    cur.execute("SELECT id, password_hash FROM users WHERE email=%s", (email,))
    user = cur.fetchone()
    if user and bcrypt.checkpw(password, user[1].encode('utf-8')):
        otp = generate_otp()
        expires_at = datetime.now() + timedelta(minutes=5)
        cur.execute("INSERT INTO otps (user_email, otp_code, expires_at, purpose) VALUES (%s, %s, %s, %s)", 
                    (email, otp, expires_at, 'login'))
        mysql.connection.commit()
        msg = Message('Your Login OTP', sender=app.config['MAIL_USERNAME'], recipients=[email])
        msg.body = f"Your OTP for login is: {otp}"
        mail.send(msg)
        session['pending_user'] = email
        session['pending_action'] = 'login'
        flash('OTP sent to your email.', 'info')
        return redirect(url_for('otp'))
    else:
        flash('Invalid email or password.', 'danger')
        return redirect(url_for('login_page'))


# REGISTER (send OTP to admin)
@app.route('/register', methods=['POST'])
def register():
    email = request.form['email']
    password = request.form['password'].encode('utf-8')
    cur = mysql.connection.cursor()
    cur.execute("SELECT id FROM users WHERE email=%s", (email,))
    if cur.fetchone():
        flash('Email already registered.', 'danger')
        return redirect(url_for('login_page'))
    otp = generate_otp()
    expires_at = datetime.now() + timedelta(minutes=5)
    cur.execute("INSERT INTO otps (user_email, otp_code, expires_at, purpose) VALUES (%s, %s, %s, %s)", 
                (email, otp, expires_at, 'register'))
    mysql.connection.commit()
    msg = Message('New User Registration OTP', sender=app.config['MAIL_USERNAME'], recipients=[ADMIN_EMAIL])
    msg.body = f"OTP for approving new user ({email}): {otp}"
    mail.send(msg)
    session['pending_user'] = email
    session['pending_password'] = password.decode()
    session['pending_action'] = 'register'
    flash('OTP sent to admin for approval. Please ask admin for the OTP.', 'info')
    return redirect(url_for('otp'))

# OTP Verification
@app.route('/otp', methods=['GET', 'POST'])
@limiter.limit("30 per hour")
def otp():
    if request.method == 'POST':
        otp = request.form['otp']
        email = session.get('pending_user')
        action = session.get('pending_action')
        cur = mysql.connection.cursor()
        cur.execute("SELECT id, expires_at FROM otps WHERE user_email=%s AND otp_code=%s AND purpose=%s ORDER BY id DESC LIMIT 1",
                    (email, otp, action))
        record = cur.fetchone()
        if record and datetime.now() < record[1]:
            if action == 'login':
                cur.execute("SELECT id, email FROM users WHERE email=%s", (email,))
                user_row = cur.fetchone()
                if user_row:
                    user = User(user_row[0], user_row[1])
                    login_user(user)
                    session['user_email'] = email
                    flash('Login successful!', 'success')
                    cur.execute("DELETE FROM otps WHERE id=%s", (record[0],))
                    mysql.connection.commit()
                    return redirect(url_for('login_page'))  # or 'index' if you want

            elif action == 'register':
                password = session.get('pending_password').encode('utf-8')
                hashed = bcrypt.hashpw(password, bcrypt.gensalt())
                cur.execute("INSERT INTO users (email, password_hash) VALUES (%s, %s)", (email, hashed.decode()))
                mysql.connection.commit()
                # Clean up OTP
                cur.execute("DELETE FROM otps WHERE id=%s", (record[0],))
                mysql.connection.commit()
                flash('Registration successful! You can now log in.', 'success')
                return redirect(url_for('login_page'))

        else:
            flash('Invalid or expired OTP.', 'danger')
    return render_template('otp.html')

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form['email']
        cur = mysql.connection.cursor()
        cur.execute("SELECT id FROM users WHERE email=%s", (email,))
        user = cur.fetchone()
        if user:
            # Generate token and expiry
            token = secrets.token_urlsafe(32)
            expires_at = datetime.now() + timedelta(hours=1)
            # Store token in DB
            cur.execute("""
                INSERT INTO password_resets (user_id, token, expires_at)
                VALUES (%s, %s, %s)
                ON DUPLICATE KEY UPDATE token=%s, expires_at=%s
            """, (user[0], token, expires_at, token, expires_at))
            mysql.connection.commit()
            reset_url = url_for('reset_password', token=token, _external=True)
            msg = Message('Password Reset Request',
                          sender=app.config['MAIL_USERNAME'],
                          recipients=[email])
            msg.body = f'Click the link to reset your password: {reset_url}\nThis link will expire in 1 hour.'
            mail.send(msg)
            flash('A password reset link has been sent to your email.', 'info')
        else:
            flash('Email not found.', 'danger')
        return redirect(url_for('forgot_password'))
    return render_template('forgot_password.html')


@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    cur = mysql.connection.cursor()
    cur.execute("SELECT user_id, expires_at FROM password_resets WHERE token=%s", (token,))
    record = cur.fetchone()
    if not record or datetime.now() > record[1]:
        flash('Invalid or expired token.', 'danger')
        return redirect(url_for('login_page'))
    if request.method == 'POST':
        password = request.form['password'].encode('utf-8')
        hashed = bcrypt.hashpw(password, bcrypt.gensalt())
        cur.execute("UPDATE users SET password_hash=%s WHERE id=%s", (hashed.decode(), record[0]))
        cur.execute("DELETE FROM password_resets WHERE token=%s", (token,))
        mysql.connection.commit()
        flash('Your password has been reset. Please log in.', 'success')
        return redirect(url_for('login_page'))
    return render_template('reset_password.html', token=token)


@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('index.html', email=session.get('user_email'))

@app.route('/logout')
@login_required
def logout():
    logout_user()
    session.clear()
    flash('Logged out.', 'info')
    return redirect(url_for('login_page'))



try:
    from docx2pdf import convert
    import pythoncom
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

from pdfrw import PdfReader, PdfWriter, PageMerge
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

docx2pdf_lock = threading.Lock()
unoconv_lock = threading.Lock()

OUTPUT_FOLDER = 'output'
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

COMPANY_TEMPLATES = {
    "ROYAL_SKY_INTERNATIONAL": 'templates/ROYAL',
    "VIVAN":'templates/VIVAN',
    "AR": 'templates/AR',
    "SNS_GLOBLE": 'templates/SNSGLOBLE'
}

SHEET_NAME = {
    "ROYAL_SKY_INTERNATIONAL": 'RS',
    "VIVAN":'VI',
    "AR": 'AR',
    "SNS_GLOBLE": 'SNS'
}


# Chnage1
COMPANY_GOOGLE_SHEETS = {
    "ROYAL_SKY_INTERNATIONAL": {
        "url": "https://docs.google.com/spreadsheets/d/1vgXggucKcJ09xXJj-mjraFnk_PH3iCEKm1iv6Teq7UI/edit?gid=787616279",
        "sheet_id": "1vgXggucKcJ09xXJj-mjraFnk_PH3iCEKm1iv6Teq7UI"
    },
    "VIVAN": {
        "url": "https://docs.google.com/spreadsheets/d/1FcU1XCAGohd_bdqO3GJIgsKoucZhJieRNM_1Jmmbf94/edit?gid=0#gid=0",
        "sheet_id": "1FcU1XCAGohd_bdqO3GJIgsKoucZhJieRNM_1Jmmbf94"
    },
    "AR": {
        "url": "https://docs.google.com/spreadsheets/d/1hYiWttZnmkma8ejd9DKEJosa_-H2jow8vsbGfNUAj3Q/edit?gid=0#gid=0",
        "sheet_id": "1hYiWttZnmkma8ejd9DKEJosa_-H2jow8vsbGfNUAj3Q"
    },
    "SNS_GLOBLE": {
        "url": "https://docs.google.com/spreadsheets/d/1vgXggucKcJ09xXJj-mjraFnk_PH3iCEKm1iv6Teq7UI/edit?gid=787616279",
        "sheet_id": "1vgXggucKcJ09xXJj-mjraFnk_PH3iCEKm1iv6Teq7UI"
    }
}


BATCH_STATUS = {}

def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            _replace_in_runs(paragraph.runs, f"{{{{{key}}}}}", str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        _replace_in_runs(paragraph.runs, f"{{{{{key}}}}}", str(value))

def _replace_in_runs(runs, placeholder, value):
    full_text = ''.join(run.text for run in runs)
    if placeholder not in full_text:
        return
    start = full_text.find(placeholder)
    while start != -1:
        end = start + len(placeholder)
        current = 0
        for run in runs:
            run_len = len(run.text)
            if current <= start < current + run_len:
                run_start = start - current
                run_end = min(run_len, end - current)
                before = run.text[:run_start]
                after = run.text[run_end:]
                run.text = before + value + after
                left = end - (current + run_len)
                if left > 0:
                    _remove_placeholder_from_next_runs(runs, runs.index(run)+1, left)
                break
            current += run_len
        full_text = ''.join(run.text for run in runs)
        start = full_text.find(placeholder)

def _remove_placeholder_from_next_runs(runs, start_idx, left):
    for i in range(start_idx, len(runs)):
        if left <= 0:
            break
        run = runs[i]
        if left >= len(run.text):
            left -= len(run.text)
            run.text = ''
        else:
            run.text = run.text[left:]
            left = 0

def fill_pdf_template(input_pdf_path, output_pdf_path, replacements):
    template_pdf = PdfReader(input_pdf_path)
    for page in template_pdf.pages:
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        y = 700
        for key, value in replacements.items():
            can.drawString(100, y, f"{key}: {value}")
            y -= 20
        can.save()
        packet.seek(0)
        overlay_pdf = PdfReader(packet)
        PageMerge(page).add(overlay_pdf.pages[0]).render()
    PdfWriter(output_pdf_path, trailer=template_pdf).write()

def convert_docx_to_pdf(docx_path, output_dir=None, timeout=30):
    if output_dir is None:
        output_dir = os.path.dirname(docx_path)

    try:
        with unoconv_lock:  # <--- Only one thread at a time can run this
            with open(docx_path, 'rb') as f:
                files = {'file': (os.path.basename(docx_path), f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                response = requests.post('http://localhost:3001/convert', files=files, timeout=timeout)
                if response.status_code != 200:
                    raise Exception("Failed to convert using unoconv API")

                pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
                pdf_path = os.path.join(output_dir, pdf_filename)
                with open(pdf_path, 'wb') as out_file:
                    out_file.write(response.content)
                return pdf_path
    except Exception as e:
        raise Exception(f"Error converting DOCX to PDF via unoconv API: {e}")

def process_documents_in_batches(passport_data_list, data, template_folder, session_output, output_format, batch_size=5):
    all_doc_files = []
    all_missing_messages = []

    for i in range(0, len(passport_data_list), batch_size):
        batch = passport_data_list[i:i+batch_size]
        max_workers = min(batch_size, len(batch))
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = [
                executor.submit(
                    generate_documents_for_passport,
                    pd,
                    data,
                    template_folder,
                    session_output,
                    output_format
                )
                for pd in batch
            ]
            for future in as_completed(futures):
                try:
                    doc_files, missing = future.result()
                    all_doc_files.extend(doc_files)
                    all_missing_messages.extend(missing)
                except Exception as e:
                    all_missing_messages.append(f"Batch error: {e}")

    return all_doc_files, all_missing_messages


def generate_documents_for_passport(passport_data, data, TEMPLATE_FOLDER, session_output, output_format):
    # print(f"START: SRNO={passport_data.get('srno')} | PASSPORTNO={passport_data.get('PASSPORTNO')}")
    try:    
        files = []
        missing_messages = []
        REQUIRED_FIELDS = {
            'request_letter': ['SPNAME', 'SPADD', 'CRNONDIDNO', 'PHONENO', 'VISAISSUEDATE'],
            'agreement':      ['SPNAME', 'SPADD', 'VISAISSUEDATE', 'PASSPORTNAME', 'PASSPORTNO' , 'JOBROLE', 'SALARY'],
            'afi_noc':        ['PASSPORTNAME', 'PASSPORTNO', 'SPNAME' , 'VISANO', 'VISAEXPIRY' , 'FEID' , 'Country Name'],
            'Annexure':       ['SPNAME' , 'Country Name','LEGAL_STATUS','AUTHORISED_SIGNATORY','ID_NO'],
            'POA_DM':         ['SPNAME', 'SPADD', 'VISAISSUEDATE', 'JOBROLE', 'SALARY']
        }

        DOC_DISPLAY_NAMES = {
            'request_letter': 'Request Letter',
            'agreement': 'Agreement',
            'afi_noc': 'Affidavit',
            'Annexure': 'Annexure',
            'POA_DM' : 'POA_DM'
        }

        sr_no = passport_data['srno']
      

        for doc_type in data.get("selectedDocs", ['agreement', 'request_letter', 'afi_noc', 'Annexure','POA_DM']):
            missing_cols = []
            for col in REQUIRED_FIELDS.get(doc_type, []):
                value = passport_data.get(col, "")
                # print(f"SRNO {passport_data.get('srno')} - {col}: '{value}'")
                if col not in passport_data.index or pd.isna(passport_data[col]) or str(passport_data[col]).strip() == "":
                    missing_cols.append(col)
                if missing_cols:
 #                       print(f"SRNO {passport_data.get('srno')} missing fields for {doc_type}: {missing_cols}")
                        continue
#                print(f"GEN: Generating {doc_type} for SRNO {passport_data.get('srno')}")
                val = passport_data[col]
                if pd.isna(val) or str(val).strip() in ('', '0', 'nan', 'NaN'):
                    missing_cols.append(col)
            if missing_cols:
                    display_name = DOC_DISPLAY_NAMES.get(doc_type, doc_type.replace('_', ' '))
                    missing_messages.append(f"{sr_no} {display_name} : {', '.join(missing_cols)}")
                    continue
            country_name = passport_data['Country Name']
            phoneno = passport_data['PHONENO']
            passport_data['PHONENO'] = phoneno

            templates_path = os.path.join(TEMPLATE_FOLDER, str(country_name))
            if not os.path.exists(templates_path):
                continue

            file_prefix = f"{sr_no} {passport_data['PASSPORTNO']}"
            # Force date format before replacing in docx
            if 'USEDATE' in passport_data and passport_data['USEDATE']:
                if isinstance(passport_data['USEDATE'], (datetime,)):
                    passport_data['USEDATE'] = passport_data['USEDATE'].strftime("%d-%m-%Y")
                else:
                    try:
                        passport_data['USEDATE'] = pd.to_datetime(passport_data['USEDATE']).strftime("%d-%m-%Y")
                    except:
                        pass

            replacements = passport_data.to_dict()
            DOC_MAP = {
                'agreement': ('agreement.docx', 'Agreement'),
                'request_letter': ('request_letter.docx', 'Request Letter'),
                'afi_noc': ('afi_noc.docx', 'Affidavit'),
                'Annexure' : ('Annexure.docx', 'Annexure'),
                'POA_DM' : ('POA_DM.docx', 'POA_DM')
            }
            template_files = [DOC_MAP[key] for key in data.get("selectedDocs", ['agreement', 'request_letter', 'afi_noc','Annexure' ,'POA_DM']) if key in DOC_MAP]

            for template_file, display_name in template_files:
                pdf_template_path = os.path.join(templates_path, template_file.replace('.docx', '.pdf'))
                output_name = f"{sr_no}-{display_name}"

                if os.path.exists(pdf_template_path):
                    output_pdf = os.path.join(session_output, f"{output_name}.pdf")
                    fill_pdf_template(pdf_template_path, output_pdf, replacements)
                    files.append({"name": f"{output_name}.pdf", "url": f"/download/{{session_id}}/{output_name}.pdf"})
                    continue

                template_path = os.path.join(templates_path, template_file)
                if not os.path.exists(template_path):
                    continue

                doc = Document(template_path)
                replace_placeholders(doc, replacements)
                output_docx = os.path.join(session_output, f"{output_name}.docx")
                doc.save(output_docx)
                #print(f"FILE WRITTEN: {output_docx}")

                if output_format == "pdf":
                    try:
                        output_pdf = convert_docx_to_pdf(output_docx, session_output)
                        #print(f"FILE WRITTEN: {output_pdf}")
                        files.append({"name": f"{output_name}.pdf", "url": f"/download/{{session_id}}/{output_name}.pdf"})
                        #print(f"Appended to files: {files[-1]}")
                    except Exception as e:
                        files.append({"name": f"{output_name}.docx", "url": f"/download/{{session_id}}/{output_name}.docx"})
                        #print(f"ERROR writing file for SRNO {passport_data.get('srno')}: {e}")
                else:
                    files.append({"name": f"{output_name}.docx", "url": f"/download/{{session_id}}/{output_name}.docx"})
                    #print(f"Appended to files: {files[-1]}")

        return files, missing_messages
    except Exception as e:
        print(f"Error generating document for passport {passport_data}: {e}")
    # Return empty list and error message so main thread can handle it gracefully

    return [], [f"Error processing passport {passport_data}: {e}"]
    # ...document generation and saving code...


def process_batch_in_background(session_id, passport_data_list, data, template_folder, session_output, output_format):
    try:
        BATCH_STATUS[session_id] = {
            "status": "processing",
            "total_batches": len(passport_data_list),
            "completed_batches": 0,
            "files": [],
            "missing_messages": []
        }
        
        # Process in small batches (5 documents per batch)
        batch_size = 5
        for i in range(0, len(passport_data_list), batch_size):
            batch = passport_data_list[i:i+batch_size]
            files, missing = process_documents_in_batches(
                batch, data, template_folder, session_output, output_format, batch_size
            )
            
            # Update status with batch results
            BATCH_STATUS[session_id]["files"].extend(files)
            BATCH_STATUS[session_id]["missing_messages"].extend(missing)
            BATCH_STATUS[session_id]["completed_batches"] += 1
            
            # Write batch results to disk for persistence
            status_path = os.path.join(session_output, "status.json")
            with open(status_path, "w") as f:
                json.dump(BATCH_STATUS[session_id], f)
        
        BATCH_STATUS[session_id]["status"] = "completed"
        # Final update to disk
        status_path = os.path.join(session_output, "status.json")
        with open(status_path, "w") as f:
            json.dump(BATCH_STATUS[session_id], f)
    
    except Exception as e:
        BATCH_STATUS[session_id]["status"] = "error"
        BATCH_STATUS[session_id]["error"] = str(e)
        status_path = os.path.join(session_output, "status.json")
        with open(status_path, "w") as f:
            json.dump(BATCH_STATUS[session_id], f)


def normalize_gsheet_dates(df, date_col):
    """
    Converts Google Sheets serial date numbers or strings to datetime.date.
    """
    # If the column is numeric (serial), convert using correct origin
    if pd.api.types.is_numeric_dtype(df[date_col]):
        # Google Sheets origin is 1899-12-30
        df[date_col] = pd.to_datetime(df[date_col], unit='d', origin='1899-12-30').dt.date
    else:
        # If already string, parse to date
        df[date_col] = pd.to_datetime(df[date_col], errors='coerce').dt.date
    return df


def parse_user_date(date_str):
    # Handles 'YYYY-MM-DD' format from HTML input[type=date]
    return datetime.strptime(date_str, "%Y-%m-%d").date()

def filter_by_date_range(df, date_col, start_date, end_date):
    return df[(df[date_col] >= start_date) & (df[date_col] <= end_date)]

def clean_crnondidno(val):
    if pd.isna(val):
        return ''
    val = str(val).strip()
    # Remove Excel/Sheets timestamps like '1970-01-09 10:16:02'
    if re.match(r"\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}", val):
        return ''
    return val

@app.route('/process', methods=['POST'])
@login_required
def process():
    try:
        data = request.get_json()
        #print('Received data:', data)
        #print('useDate:', data.get('useDate'))
        passport_number = data.get("passportNumber")

        company = data.get("company")
        TEMPLATE_FOLDER = COMPANY_TEMPLATES.get(company)
        SHEET = SHEET_NAME.get(company)

        if not TEMPLATE_FOLDER or not os.path.exists(TEMPLATE_FOLDER):
            return jsonify({"success": False, "message": "Invalid company selected"})


        # Chnage2 start
        company_sheet = COMPANY_GOOGLE_SHEETS.get(company)
        if not company_sheet:
            return jsonify({"success": False, "message": "Company sheet configuration not found"})
        
        google_sheet_url = company_sheet["url"]
        sheet_id = company_sheet["sheet_id"]
        
        # Build CSV URL with company-specific sheet ID
        csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={SHEET}"
        # Chnage 2 end

        response = requests.get(csv_url)
        response.raise_for_status()

        # ==== ADD THIS BLOCK: Robust DataFrame reading and column handling ====
        df = pd.read_csv(io.StringIO(response.text), dtype=str, skip_blank_lines=True)
        df.columns = [col.strip() for col in df.columns]  # Clean up column names

        df_all = df.copy()  # <-- Define df_all here, before using it

        df['Country Name'] = df['Country Name'].astype(str).str.strip()
        df_all['Country Name'] = df_all['Country Name'].astype(str).str.strip()
        # Optionally drop Unnamed columns (extra empty columns from sheet)
        df = df.loc[:, ~df.columns.str.contains('^Unnamed')]
        df['CRNONDIDNO'] = df['CRNONDIDNO'].apply(clean_crnondidno)
        df_all['CRNONDIDNO'] = df_all['CRNONDIDNO'].apply(clean_crnondidno)

        # If you have required columns, ensure they exist
        required_columns = [
            'DATE', 'srno', 'USEDATE', 'CODE', 'PASSPORTNAME', 'SPNAME',  # Add more as needed
        ]
        for col in required_columns:
            if col not in df.columns:
                df[col] = ''
        # ==== END BLOCK ====

        df = df[df['PASSPORTNO'].notnull() & (df['PASSPORTNO'] != '')]

        # --- NEW: Load the full sheet for passport search ---
        df_all = df.copy()

        # Normalize and parse the date correctly
        use_date_str = data.get('useDate')
        if not use_date_str or not isinstance(use_date_str, str) or not use_date_str.strip():
            return jsonify({"success": False, "message": "No date selected. Please select a date."})

        try:
            if '-' in use_date_str:
                if use_date_str.count('-') == 2:
                    parts = use_date_str.split('-')
                    if len(parts[0]) == 4:
                        use_date = datetime.strptime(use_date_str, "%Y-%m-%d").date()
                    else:
                        use_date = datetime.strptime(use_date_str, "%d-%m-%Y").date()
                else:
                    raise ValueError("Invalid date format")
            else:
                raise ValueError("Invalid date format")
        except Exception as e:
            return jsonify({"success": False, "message": f"Invalid date format selected: {e}"})

        # Convert the sheet column to datetime objects with the correct format
        df['USEDATE'] = pd.to_datetime(df['USEDATE'], format='%d-%m-%Y', errors='coerce').dt.date
        df = df.dropna(subset=['USEDATE'])  # Remove rows with invalid dates

        lines = response.text.splitlines()
        header_idx = None
        for i, line in enumerate(lines[:5]):
            if any('PASSPORTNO' in col.replace(' ', '').upper() for col in line.split(',')):
                header_idx = i
                break
        if header_idx is None:
            return jsonify({"success": False, "message": "PASSPORTNO column missing in sheet."})

        # Date filtering
        df = df[df['USEDATE'] == use_date]  # Filter after converting to datetime

        # SRNO/Passport filtering
        selected_rows = pd.DataFrame()
        if data.get("startSrno") and data.get("endSrno"):
            start = int(data["startSrno"])
            end = int(data["endSrno"])
            # Clean whitespace and ensure valid integer conversion
            df['srno_clean'] = df['srno'].astype(str).str.strip()
            df = df[df['srno_clean'].str.isdigit()]
            df['srno_int'] = df['srno_clean'].astype(int)
            df = df[(df['srno_int'] >= start) & (df['srno_int'] <= end)]
            #print("Filtered SRNOs:", df['srno_int'].tolist())
            #print(df[['srno', 'srno_clean', 'srno_int']])
            selected_rows = df
        elif passport_number:
            # --- Use df_all for passport number search (ignore date filter) ---
            passport_search = str(passport_number).strip().upper()
            df_all['PASSPORTNO_CLEAN'] = df_all['PASSPORTNO'].astype(str).str.strip().str.upper()
            #print("Searching for passport:", passport_search)
            #print("Available passport numbers:", df_all['PASSPORTNO_CLEAN'].unique())

            # Find all rows for this passport
            passport_rows = df_all[df_all['PASSPORTNO_CLEAN'] == passport_search]
            #print("Rows found for passport:", passport_rows.shape[0])

            if passport_rows.empty:
                # Try fuzzy match for debugging
                fuzzy_matches = df_all[df_all['PASSPORTNO_CLEAN'].str.contains(passport_search, na=False)]
                #print("Fuzzy matches found:", fuzzy_matches['PASSPORTNO_CLEAN'].tolist())
                return jsonify({"success": False, "message": "Passport number not found."})

            # Convert USEDATE to datetime for sorting (handle errors gracefully)
            passport_rows['USEDATE_DT'] = pd.to_datetime(passport_rows['USEDATE'], format='%d-%m-%Y', errors='coerce')
            passport_rows = passport_rows.dropna(subset=['USEDATE_DT'])
            if passport_rows.empty:
                return jsonify({"success": False, "message": "No valid date found for this passport number."})

            # Get the row with the latest USEDATE
            latest_row = passport_rows.sort_values('USEDATE_DT', ascending=False).head(1)
            selected_rows = latest_row

        else:
            return jsonify({"success": False, "message": "Invalid filter parameters"})

        if selected_rows.empty:
            return jsonify({"success": False, "message": "No matching records found"})

        # Session setup
        session_id = str(uuid.uuid4())
        session_output = os.path.join(OUTPUT_FOLDER, session_id)
        os.makedirs(session_output, exist_ok=True)

        # Start background processing
        passport_data_list = [row.copy() for _, row in selected_rows.iterrows()]
        Thread(
            target=process_batch_in_background,
            args=(session_id, passport_data_list, data, TEMPLATE_FOLDER, session_output, data.get("outputFormat", "pdf"))
        ).start()

        return jsonify({
            "success": True,
            "session_id": session_id,
            "message": "Document generation started"
        })

    except ValueError as e:
        return jsonify({"success": False, "message": f"Date format error: {str(e)}"})
    except Exception as e:
        return jsonify({"success": False, "message": f"Processing error: {str(e)}"})


@app.route('/batch-status/<session_id>', methods=['GET'])
@login_required
def batch_status(session_id):
    try:
        # Try to get from memory first
        if session_id in BATCH_STATUS:
            status = BATCH_STATUS[session_id].copy()

        else:
            # Try to load from disk
            status_path = os.path.join(OUTPUT_FOLDER, session_id, "status.json")
            if os.path.exists(status_path):
                with open(status_path, "r") as f:
                    status = json.load(f)
            else:
                return jsonify({"success": False, "message": "Session not found"})
        
        # Update file URLs with session_id
        for file in status.get("files", []):
            if "{session_id}" in file.get("url", ""):
                file["url"] = file["url"].replace("{session_id}", session_id)
        
        return jsonify({
            "success": True,
            "status": status["status"],
            "total_batches": status.get("total_batches", 0),
            "completed_batches": status.get("completed_batches", 0),
            "files": status.get("files", []),
            "missing_values": status.get("missing_messages", []),
            "error": status.get("error")
        })
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})


@app.route('/search-srnos', methods=['POST'])
@login_required
def search_srnos():
    data = request.get_json()
    use_date = data.get('useDate')
    company = data.get('company')
    SHEET = SHEET_NAME.get(company)
    

    # chnages3
    company_sheet = COMPANY_GOOGLE_SHEETS.get(company)
    if not company_sheet:
        return jsonify({"success": False, "message": "Company sheet configuration not found"})
    
    sheet_id = company_sheet["sheet_id"]
    csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={SHEET}"
#    chnage 3 end
    response = requests.get(csv_url)
    response.raise_for_status()

    df = pd.read_csv(io.StringIO(response.text), dtype=str, skip_blank_lines=True)
    # Parse dates
    df['USEDATE'] = pd.to_datetime(df['USEDATE'], errors='coerce', dayfirst=True).dt.date
    input_date = pd.to_datetime(use_date, errors='coerce', dayfirst=False).date()
    filtered = df[df['USEDATE'] == input_date]
    srnos = filtered['srno'].dropna().unique().tolist()
    return jsonify({
        "count": len(srnos),
        "srnos": srnos
    })



@app.route('/download/<session_id>/<filename>')
@login_required
def download(session_id, filename):
    file_path = os.path.join(OUTPUT_FOLDER, session_id, filename)
    session_folder = os.path.join(OUTPUT_FOLDER, session_id)

    return send_from_directory(session_folder, filename, as_attachment=True)

@app.route('/download-all', methods=['GET', 'POST'])
@login_required
def download_all():
    if request.method == 'POST':
        data = request.get_json() or {}
        session_id = data.get('session_id')
        file_prefix = data.get('file_prefix')
    else:
        session_id = request.args.get('session_id')
        file_prefix = request.args.get('file_prefix')

    if not session_id:
        return "No files to download. Generate documents first.", 404

    session_dir = os.path.join(OUTPUT_FOLDER, session_id)
    if not os.path.exists(session_dir):
        return "Session files not found", 404

    memory_file = io.BytesIO()
    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file in os.listdir(session_dir):
            file_path = os.path.join(session_dir, file)
            if os.path.isfile(file_path) and file.lower().endswith('.pdf'):
                zipf.write(file_path, arcname=file)

    memory_file.seek(0)
    zip_name = f"{file_prefix}.zip" if file_prefix else "all_documents.zip"


    return send_file(memory_file, mimetype='application/zip', as_attachment=True, download_name=zip_name)

@app.route('/filter-by-date', methods=['POST'])
@login_required
def filter_by_date():
    try:
        data = request.get_json()
        selected_date = data.get('date')
        company = data.get('company')

        if not selected_date or not company:
            return jsonify({"success": False, "message": "Missing date or company"})

        TEMPLATE_FOLDER = COMPANY_TEMPLATES.get(company)
        SHEET = SHEET_NAME.get(company)

        # chnage4
        company_sheet = COMPANY_GOOGLE_SHEETS.get(company)
        if not company_sheet:
            return jsonify({"success": False, "message": "Company sheet configuration not found"})
        
        sheet_id = company_sheet["sheet_id"]
        csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={SHEET}"
        # chnage 4 end 

        df = pd.read_csv(csv_url, dtype=str)

        if 'USEDATE' not in df.columns or 'srno' not in df.columns:
            return jsonify({"success": False, "message": "Required columns missing."})

        match_date = datetime.strptime(selected_date, '%Y-%m-%d').strftime('%d-%m-%Y')
        df['USEDATE'] = df['USEDATE'].astype(str).str.strip()
        df['srno'] = df['srno'].astype(str).str.strip()

        filtered = df[(df['USEDATE'] == match_date) & (df['srno'] != '0')]

        srnos = []
        for val in filtered['srno']:
            try:
                srnos.append(int(val))
            except:
                continue

        return jsonify({
            "success": True,
            "total": len(srnos),
            "srnos": srnos
        })

    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

if __name__ == '__main__':
    app.run(debug=True)
